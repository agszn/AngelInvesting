from django.shortcuts import render
from django.contrib.auth.decorators import login_required

from django.shortcuts import render, redirect
from .models import *
from django.contrib import messages
from django.db.models import Q

def BlogSM(request):
    blogs = Blog.objects.all().order_by('-date')  # default: show all

    # === FILTER LOGIC ===
    filter_date = request.GET.get('filter_date')
    filter_heading = request.GET.get('filter_heading')

    if filter_date:
        blogs = blogs.filter(date=filter_date)

    if filter_heading:
        blogs = blogs.filter(heading__icontains=filter_heading)

    # === CREATE BLOG ===
    if request.method == 'POST':
        banner = request.FILES.get('banner')
        date = request.POST.get('date')
        heading = request.POST.get('heading')
        description = request.POST.get('description')

        blog = Blog(
            date=date,
            heading=heading,
            description=description,
        )
        if banner:
            blog.banner = banner
        blog.save()
        messages.success(request, "Blog created successfully!")
        return redirect('SM_User:BlogSM')  # replace with your URL name

    context = {
        'blogs': blogs,
        'filter_date': filter_date or '',
        'filter_heading': filter_heading or '',
    }

    return render(request, 'blogs/BlogSM.html', context)

from django.shortcuts import render, get_object_or_404
from .models import Blog

def blog_detail_view_SM(request, blog_id):
    blog = get_object_or_404(Blog, pk=blog_id)
    related_blogs = Blog.objects.exclude(id=blog_id).order_by('-date')[:3]
    return render(request, 'blogs/blog_detailST.html', {'blog': blog, 'related_blogs': related_blogs})


from django.shortcuts import render, redirect, get_object_or_404
from .models import Blog
from .forms import BlogForm
from django.contrib import messages

def create_blog(request):
    if request.method == 'POST':
        form = BlogForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, "Blog created successfully!")
            return redirect('SM_User:BlogSM')
    else:
        form = BlogForm()
    return render(request, 'blogs/blog_form.html', {'form': form, 'title': 'Create Blog'})

def edit_blog(request, blog_id):
    blog = get_object_or_404(Blog, id=blog_id)
    if request.method == 'POST':
        form = BlogForm(request.POST, request.FILES, instance=blog)
        if form.is_valid():
            form.save()
            messages.success(request, "Blog updated successfully!")
            return redirect('SM_User:BlogSM')
    else:
        form = BlogForm(instance=blog)
    return render(request, 'blogs/blog_form.html', {'form': form, 'title': 'Edit Blog'})

def delete_blog(request, blog_id):
    blog = get_object_or_404(Blog, id=blog_id)
    if request.method == 'POST':
        blog.delete()
        messages.success(request, "Blog deleted successfully!")
        return redirect('SM_User:BlogSM')
    return redirect('SM_User:BlogSM')

# views.py

from django.shortcuts import render, redirect
from .models import HeroSectionBanner
from django.core.files.storage import default_storage
from django.shortcuts import render, redirect
from .models import *
from unlisted_stock_marketplace.models import *

@login_required
def HomepageBannerSM(request):
    banner = HeroSectionBanner.objects.filter(is_active=True).first()

    if request.method == 'POST' and request.FILES.get('banner'):
        # Deactivate previous banners
        HeroSectionBanner.objects.filter(is_active=True).update(is_active=False)
        
        new_banner = HeroSectionBanner.objects.create(
            image=request.FILES['banner'],
            is_active=True
        )
        return redirect('SM_User:HomepageBannerSM')  # refresh page

    return render(request, 'HomepageBannerSM.html', {'banner': banner})

from django.views.decorators.http import require_POST
from django.http import HttpResponseRedirect
from django.urls import reverse

@require_POST
def delete_banner(request):
    HeroSectionBanner.objects.filter(is_active=True).delete()
    return redirect('SM_User:HomepageBannerSM') 

from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.db.models import Q
from django.contrib import messages
from decimal import Decimal, InvalidOperation
import csv, io
from django.http import HttpResponse

from decimal import Decimal, InvalidOperation
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.db.models import Q
from .forms import *


# @login_required
# def UnlistedStocksUpdateSM(request):
#     today = timezone.now().date()

#     if request.method == "POST":
#         stock_id = request.POST.get('stock_id')
#         share_price_str = request.POST.get('share_price')
#         conviction_level = request.POST.get('conviction_level')
#         lot_size_str = request.POST.get('lot_size')

#         if stock_id:
#             stock = get_object_or_404(StockData, pk=stock_id)

#             # Only update lot_size directly in StockData
#             try:
#                 lot_size = int(lot_size_str) if lot_size_str else None
#                 if lot_size is not None:
#                     stock.lot = lot_size
#                     stock.save()
#             except (ValueError, TypeError):
#                 pass  # Ignore invalid lot sizes

#             # Create or update the snapshot for today
#             snapshot, _ = StockDailySnapshot.objects.get_or_create(stock=stock, date=today)

#             try:
#                 share_price = Decimal(share_price_str) if share_price_str else None
#                 if share_price is not None:
#                     snapshot.share_price = share_price
#             except (InvalidOperation, TypeError):
#                 pass

#             if conviction_level:
#                 snapshot.conviction_level = conviction_level

#             snapshot.save(update_stockdata=False)

#             return redirect('SM_User:UnlistedStocksUpdateSM')

#     # Handle GET request
#     query = request.GET.get('q', '').strip()
#     all_stocks = StockData.objects.all()

#     if query:
#         all_stocks = all_stocks.filter(Q(company_name__icontains=query) | Q(sector__icontains=query))

#     snapshots = []
#     for stock in all_stocks:
#         snapshot = StockDailySnapshot.objects.filter(stock=stock, date=today).first()
#         if snapshot:
#             snapshots.append(snapshot)
#         else:
#             # Fallback to the most recent snapshot, or a placeholder with today's date
#             latest_snapshot = StockDailySnapshot.objects.filter(stock=stock).order_by('-date').first()
#             snapshots.append(latest_snapshot or StockDailySnapshot(stock=stock, date=today))

#     # Optional: Display CSV upload result feedback
#     upload_result = request.session.pop('csv_update_result', None)

#     return render(request, 'UnlistedStocksUpdateSM.html', {
#         'snapshots': snapshots,
#         'conviction_choices': CONVICTION_CHOICES,
#         'search_query': query,
#         'today': today,
#         'upload_result': upload_result,
#     })


from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from decimal import Decimal, InvalidOperation
from django.db.models import Q
@login_required
def UnlistedStocksUpdateSM(request):
    today = timezone.now().date()

    if request.method == "POST":
        stock_id = request.POST.get('stock_id')
        share_price_str = request.POST.get('share_price')
        ltp_str = request.POST.get('ltp')
        conviction_level = request.POST.get('conviction_level')
        lot_size_str = request.POST.get('lot_size')
        company_name = request.POST.get('company_name')
        sector = request.POST.get('sector')

        if stock_id:
            stock = get_object_or_404(StockData, pk=stock_id)

            # Update editable stock fields
            if company_name:
                stock.company_name = company_name
            if sector:
                stock.sector = sector
            try:
                lot_size = int(lot_size_str) if lot_size_str else None
                if lot_size is not None:
                    stock.lot = lot_size
            except (ValueError, TypeError):
                pass

            # ✅ Avoid unnecessary history logging
            stock.save(suppress_history=True)

            # Create or update snapshot for today
            snapshot, _ = StockDailySnapshot.objects.get_or_create(stock=stock, date=today)

            try:
                ltp = Decimal(ltp_str) if ltp_str else None
                if ltp is not None:
                    snapshot.ltp = ltp
            except (InvalidOperation, TypeError):
                pass

            try:
                share_price = Decimal(share_price_str) if share_price_str else None
                if share_price is not None:
                    snapshot.share_price = share_price
            except (InvalidOperation, TypeError):
                pass

            if conviction_level:
                snapshot.conviction_level = conviction_level

            snapshot.save()  # no need for update_stockdata=False

        return redirect('SM_User:UnlistedStocksUpdateSM') # ✅ RETURN FIXED

    # --- GET request handling ---
    query = request.GET.get('q', '').strip()
    all_stocks = StockData.objects.all()
    if query:
        all_stocks = all_stocks.filter(Q(company_name__icontains=query) | Q(sector__icontains=query))

    snapshots = []
    for stock in all_stocks:
        snapshot = StockDailySnapshot.objects.filter(stock=stock, date=today).first()
        if snapshot:
            snapshots.append(snapshot)
        else:
            latest_snapshot = StockDailySnapshot.objects.filter(stock=stock).order_by('-date').first()
            snapshots.append(latest_snapshot or StockDailySnapshot(stock=stock, date=today))

    upload_result = request.session.pop('csv_update_result', None)

    return render(request, 'UnlistedStocksUpdateSM.html', {
        'snapshots': snapshots,
        'conviction_choices': CONVICTION_CHOICES,
        'search_query': query,
        'today': today,
        'upload_result': upload_result,
    })


from django.db.models import Max
@login_required
def download_unlisted_stocks_csv(request):
    # Get the latest snapshot date per stock
    latest_dates = (
        StockDailySnapshot.objects
        .values('stock')
        .annotate(latest_date=Max('date'))
    )

    # Build a dictionary {stock_id: latest_date}
    latest_date_map = {item['stock']: item['latest_date'] for item in latest_dates}

    # Query snapshots matching the latest date per stock
    latest_snapshots = StockDailySnapshot.objects.filter(
        *[
            # For each stock, filter snapshot on stock=stock_id and date=latest_date
            # But Django ORM can’t do this in a single query without subqueries
            # So do a filter with Q or subquery here:
            # We'll use a Q expression below
        ]
    )

    # The above approach is tricky in one query, so simpler approach:
    # Get all latest snapshots with a subquery:

    from django.db.models import OuterRef, Subquery

    latest_snapshot_subquery = StockDailySnapshot.objects.filter(
        stock=OuterRef('stock')
    ).order_by('-date').values('pk')[:1]

    latest_snapshots = StockDailySnapshot.objects.filter(
        pk__in=Subquery(latest_snapshot_subquery)
    ).select_related('stock')

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="unlisted_stocks_update.csv"'

    writer = csv.writer(response)
    writer.writerow(['Company Name', 'Conviction Level', 'Share Price'])

    for snapshot in latest_snapshots:
        # Filter only unlisted stocks if needed:
        if snapshot.stock.stock_type == 'Unlisted':
            writer.writerow([snapshot.stock.company_name, snapshot.conviction_level, snapshot.share_price])

    return response

@login_required
def upload_unlisted_stocks_csv(request):
    if request.method == "POST" and request.FILES.get("csv_file"):
        csv_file = request.FILES['csv_file']
        decoded_file = csv_file.read().decode('utf-8')
        io_string = io.StringIO(decoded_file)
        reader = csv.DictReader(io_string)

        updated = 0
        skipped_names = []
        updated_ids = []
        today = timezone.now().date()

        for row in reader:
            name = row.get('Company Name')
            conviction = row.get('Conviction Level')
            price = row.get('Share Price')

            if not name:
                continue

            stock = StockData.objects.filter(company_name__iexact=name.strip()).first()
            if not stock:
                skipped_names.append(name.strip())
                continue

            try:
                price_decimal = Decimal(price.strip()) if price else None
            except (InvalidOperation, TypeError):
                skipped_names.append(name.strip())
                continue

            snapshot, _ = StockDailySnapshot.objects.get_or_create(stock=stock, date=today)

            if conviction:
                snapshot.conviction_level = conviction.strip()
            if price_decimal is not None:
                snapshot.share_price = price_decimal

            snapshot.save()
            updated_ids.append(snapshot.stock.id)
            updated += 1

        request.session['csv_update_result'] = {
            'updated_ids': updated_ids,
            'skipped_names': skipped_names,
            'updated_count': updated
        }

        return redirect('SM_User:UnlistedStocksUpdateSM')

    return redirect('SM_User:UnlistedStocksUpdateSM')


# update excel
import pandas as pd
import re
from decimal import Decimal, InvalidOperation
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from django.shortcuts import redirect
from unlisted_stock_marketplace.models import StockData, StockDailySnapshot
import pandas as pd
import re
from decimal import Decimal, InvalidOperation
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from django.shortcuts import redirect

import openpyxl
import openpyxl
from decimal import Decimal, InvalidOperation
from django.shortcuts import render, redirect
from django.utils import timezone
from django.contrib import messages

from django.views.decorators.csrf import csrf_exempt

import openpyxl
from decimal import Decimal, InvalidOperation
from django.shortcuts import render, redirect
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
import openpyxl
from decimal import Decimal, InvalidOperation
from django.shortcuts import render, redirect
from django.contrib import messages
from django.utils import timezone

from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from difflib import get_close_matches
from decimal import Decimal, InvalidOperation
from difflib import get_close_matches
import openpyxl
from django.shortcuts import render, redirect
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required




# @login_required
# @csrf_exempt
# def upload_unlisted_stocks_excel(request):
#     if request.method == 'POST' and request.FILES.get('excel_file'):
#         excel_file = request.FILES['excel_file']
#         wb = openpyxl.load_workbook(excel_file)
#         sheet = wb.active

#         today = timezone.now().date()
#         yesterday = today - timezone.timedelta(days=1)

#         updated_ids = []
#         skipped_names = []
#         failed_updates = []
#         newly_added = []
#         fallback_ids = []  # for rows with 0 fallback

#         for idx, row in enumerate(sheet.iter_rows(min_row=2), start=2):
#             try:
#                 company_name = str(row[1].value).strip() if row[1].value else None
#                 raw_conviction = str(row[2].value).strip() if row[2].value else None
#                 conviction_level = raw_conviction.title() if raw_conviction else None
#                 price_today = row[3].value
#                 price_yesterday = row[4].value
#                 lot_size = row[5].value

#                 if not company_name:
#                     continue

#                 stock = StockData.objects.filter(company_name__iexact=company_name).first()
#                 if not stock:
#                     stock = StockData.objects.filter(scrip_name__iexact=company_name).first()

#                 if not stock:
#                     # Suggest similar names
#                     all_names = list(StockData.objects.values_list('company_name', flat=True)) + \
#                                 list(StockData.objects.values_list('scrip_name', flat=True))
#                     close = get_close_matches(company_name, all_names, n=3, cutoff=0.8)
#                     if close:
#                         skipped_names.append(company_name)
#                         continue

#                     # Create new StockData
#                     stock = StockData.objects.create(
#                         company_name=company_name,
#                         scrip_name=company_name,
#                         lot=int(lot_size) if lot_size else 100
#                     )
#                     newly_added.append(company_name)

#                 # Create/update yesterday snapshot
#                 try:
#                     ltp = Decimal(price_yesterday) if price_yesterday else Decimal('0.00')
#                 except (InvalidOperation, TypeError):
#                     ltp = Decimal('0.00')
#                     fallback_ids.append(str(stock.id))

#                 StockDailySnapshot.objects.get_or_create(
#                     stock=stock,
#                     date=yesterday,
#                     defaults={'share_price': ltp}
#                 )

#                 # Create or update today's snapshot
#                 snapshot, _ = StockDailySnapshot.objects.get_or_create(stock=stock, date=today)

#                 if conviction_level:
#                     if not snapshot.conviction_level or snapshot.conviction_level != conviction_level:
#                         snapshot.conviction_level = conviction_level

#                 try:
#                     share_price = Decimal(price_today) if price_today else Decimal('0.00')
#                 except (InvalidOperation, TypeError):
#                     share_price = Decimal('0.00')
#                     fallback_ids.append(str(stock.id))

#                 snapshot.share_price = share_price

#                 try:
#                     lot = int(lot_size) if lot_size else 0
#                 except (ValueError, TypeError):
#                     lot = 0
#                     fallback_ids.append(str(stock.id))

#                 snapshot.lot = lot

#                 snapshot.save()
#                 updated_ids.append(str(stock.id))

#             except Exception as e:
#                 failed_updates.append(f"{company_name}: {str(e)}")

#         context = {
#             'upload_result': {
#                 'updated_ids': updated_ids,
#                 'skipped_names': skipped_names,
#                 'failed_updates': failed_updates,
#                 'newly_added': newly_added,
#                 'fallback_ids': fallback_ids,  # for highlighting fallback 0 values
#             }
#         }
#         return render(request, 'UnlistedStocksUpdateSM.html', context)

#     return redirect('SM_User:UnlistedStocksUpdateSM')




from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from decimal import Decimal, InvalidOperation
import datetime
import openpyxl
from difflib import get_close_matches
from openpyxl.utils.datetime import from_excel

from unlisted_stock_marketplace.models import StockData, StockDailySnapshot


@login_required
@csrf_exempt
def upload_unlisted_stocks_excel(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        wb = openpyxl.load_workbook(excel_file)
        sheet = wb.active

        # ✅ Extract snapshot date from D2 (row 2, column 4)
        header_date_raw = sheet.cell(row=2, column=4).value  # Cell D2
        if isinstance(header_date_raw, datetime.datetime):
            snapshot_date = header_date_raw.date()
        elif isinstance(header_date_raw, datetime.date):
            snapshot_date = header_date_raw
        elif isinstance(header_date_raw, (int, float)):
            snapshot_date = from_excel(header_date_raw)
        else:
            snapshot_date = timezone.now().date()

        updated_ids = []
        skipped_names = []
        failed_updates = []
        newly_added = []
        fallback_ids = []

        # Start from row 3, assuming headers in row 1 and 2
        for idx, row in enumerate(sheet.iter_rows(min_row=3), start=3):
            try:
                company_name = str(row[1].value).strip() if row[1].value else None
                raw_conviction = str(row[2].value).strip() if row[2].value else None
                conviction_level = raw_conviction.title() if raw_conviction else None
                price_today = row[3].value  # ✅ Column D
                lot_size = row[5].value     # ✅ Column F

                if not company_name:
                    continue

                # --- Find or create stock ---
                stock = StockData.objects.filter(company_name__iexact=company_name).first()
                if not stock:
                    stock = StockData.objects.filter(scrip_name__iexact=company_name).first()

                if not stock:
                    # Try to suggest names
                    all_names = list(StockData.objects.values_list('company_name', flat=True)) + \
                                list(StockData.objects.values_list('scrip_name', flat=True))
                    close = get_close_matches(company_name, all_names, n=3, cutoff=0.8)
                    if close:
                        skipped_names.append(company_name)
                        continue

                    # Create new StockData
                    stock = StockData.objects.create(
                        company_name=company_name,
                        scrip_name=company_name,
                        lot=int(lot_size) if isinstance(lot_size, (int, float)) else 100
                    )
                    newly_added.append(company_name)

                # --- Create or update snapshot for extracted date ---
                snapshot, created = StockDailySnapshot.objects.get_or_create(stock=stock, date=snapshot_date)

                # --- Handle share price with fallback ---
                try:
                    share_price = Decimal(price_today)
                except (InvalidOperation, TypeError, ValueError):
                    share_price = stock.share_price if stock.share_price else Decimal('0.00')
                    fallback_ids.append(str(stock.id))

                snapshot.share_price = share_price

                # --- Handle conviction level ---
                if conviction_level:
                    snapshot.conviction_level = conviction_level

                # --- Handle lot size with fallback ---
                try:
                    lot = int(lot_size) if lot_size else (stock.lot or 0)
                except (ValueError, TypeError):
                    lot = stock.lot if stock.lot else 0
                    fallback_ids.append(str(stock.id))

                snapshot.lot = lot

                # --- Save to trigger auto ltp + sync ---
                snapshot.save()

                updated_ids.append(str(stock.id))

            except Exception as e:
                failed_updates.append(f"{company_name or 'Unknown'} (row {idx}): {str(e)}")

        context = {
            'upload_result': {
                'updated_ids': updated_ids,
                'skipped_names': skipped_names,
                'failed_updates': failed_updates,
                'newly_added': newly_added,
                'fallback_ids': fallback_ids,
            }
        }

        return render(request, 'UnlistedStocksUpdateSM.html', context)

    return redirect('SM_User:UnlistedStocksUpdateSM')



# custom fields
from django.shortcuts import render, get_object_or_404, redirect
from unlisted_stock_marketplace.models import *
from .forms import *

@login_required
def custom_field_list(request):
    definitions = CustomFieldDefinition.objects.select_related('stock').all()
    return render(request, 'custom_fields/definition_list.html', {'definitions': definitions})

@login_required
def custom_field_create(request):
    if request.method == 'POST':
        form = CustomFieldDefinitionForm(request.POST)
        formset = CustomFieldValueFormSet(request.POST)

        if form.is_valid() and formset.is_valid():
            definition = form.save()
            formset.instance = definition
            formset.save()
            return redirect('SM_User:custom_field_list')
    else:
        form = CustomFieldDefinitionForm()
        formset = CustomFieldValueFormSet()

    return render(request, 'custom_fields/definition_form.html', {'form': form, 'formset': formset})

@login_required
def custom_field_edit(request, pk):
    definition = get_object_or_404(CustomFieldDefinition, pk=pk)
    if request.method == 'POST':
        form = CustomFieldDefinitionForm(request.POST, instance=definition)
        formset = CustomFieldValueFormSet(request.POST, instance=definition)

        if form.is_valid() and formset.is_valid():
            form.save()
            formset.save()
            return redirect('SM_User:custom_field_list')
    else:
        form = CustomFieldDefinitionForm(instance=definition)
        formset = CustomFieldValueFormSet(instance=definition)

    return render(request, 'custom_fields/definition_form.html', {'form': form, 'formset': formset})


# add broker
from django.shortcuts import render, get_object_or_404, redirect
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Broker, Advisor
from .forms import BrokerForm, AdvisorForm  # You’ll need to create these forms

# ----------- BROKER VIEWS -----------

class BrokerListView(ListView):
    model = Broker
    template_name = 'broker/broker_list.html'
    context_object_name = 'brokers'

class BrokerCreateView(CreateView):
    model = Broker
    form_class = BrokerForm
    template_name = 'broker/broker_form.html'
    success_url = reverse_lazy('SM_User:broker-list')

class BrokerUpdateView(UpdateView):
    model = Broker
    form_class = BrokerForm
    template_name = 'broker/broker_form.html'
    success_url = reverse_lazy('SM_User:broker-list')

class BrokerDeleteView(DeleteView):
    model = Broker
    template_name = 'broker/broker_confirm_delete.html'
    success_url = reverse_lazy('SM_User:broker-list')

# ----------- ADVISOR VIEWS -----------

# views.py
class AdvisorTypeListView(ListView):
    model = Advisor
    template_name = 'advisor/advisor_type_list.html'
    context_object_name = 'types'

class AdvisorTypeCreateView(CreateView):
    model = Advisor
    fields = ['advisor_type']
    template_name = 'advisor/advisor_type_form.html'
    success_url = reverse_lazy('SM_User:advisor-type-list')

from django.views.generic import DeleteView
from django.urls import reverse_lazy

class AdvisorTypeDeleteView(DeleteView):
    model = Advisor
    template_name = 'advisor/advisor_type_confirm_delete.html'
    success_url = reverse_lazy('SM_User:advisor-type-list')



# stockData
from django.shortcuts import render, get_object_or_404, redirect
from .forms import StockDataForm
from django.db.models import Q

def stockdata_crud(request):
    search_query = request.GET.get('search', '')
    sector_filter = request.GET.get('sector', '')
    category_filter = request.GET.get('category', '')
    conviction_filter = request.GET.get('conviction_level', '')
    stock_type_filter = request.GET.get('stock_type', '')

    queryset = StockData.objects.all()

    if search_query:
        queryset = queryset.filter(
            Q(company_name__icontains=search_query) |
            Q(scrip_name__icontains=search_query) |
            Q(isin_no__icontains=search_query)
        )

    if sector_filter:
        queryset = queryset.filter(sector=sector_filter)

    if category_filter:
        queryset = queryset.filter(category=category_filter)

    if conviction_filter:
        queryset = queryset.filter(conviction_level=conviction_filter)

    if stock_type_filter:
        queryset = queryset.filter(stock_type=stock_type_filter)

    sectors = StockData.objects.values_list('sector', flat=True).distinct()
    categories = StockData.objects.values_list('category', flat=True).distinct()
    conviction_levels = StockData.objects.values_list('conviction_level', flat=True).distinct()
    stock_types = StockData.objects.values_list('stock_type', flat=True).distinct()

    form = StockDataForm()

    return render(request, 'StockDataTemplate.html', {
        'stocks': queryset,
        'form': form,
        'filters': {
            'search_query': search_query,
            'sector': sector_filter,
            'category': category_filter,
            'conviction_level': conviction_filter,
            'stock_type': stock_type_filter,
        },
        'sectors': sectors,
        'categories': categories,
        'conviction_levels': conviction_levels,
        'stock_types': stock_types,
    })

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, get_object_or_404, redirect
from .forms import (
    StockDataForm,
    ReportFormSet,
    ShareholdingPatternFormSet,
    CompanyRelationFormSet,
    PrincipalBusinessActivityFormSet
)

# @csrf_exempt
# def edit_stockdata(request, pk):
#     stock = get_object_or_404(StockData, pk=pk)

#     if request.method == 'POST':
#         form = StockDataForm(request.POST, request.FILES, instance=stock)
#         report_formset = ReportFormSet(request.POST, instance=stock, prefix='report')
#         shareholding_formset = ShareholdingPatternFormSet(request.POST, instance=stock, prefix='shareholding')
#         relation_formset = CompanyRelationFormSet(request.POST, instance=stock, prefix='relation')
#         business_formset = PrincipalBusinessActivityFormSet(request.POST, instance=stock, prefix='business')

#         if all([
#             form.is_valid(),
#             report_formset.is_valid(),
#             shareholding_formset.is_valid(),
#             relation_formset.is_valid(),
#             business_formset.is_valid()
#         ]):
#             form.save()
#             report_formset.save()
#             shareholding_formset.save()
#             relation_formset.save()
#             business_formset.save()
#             return redirect('SM_User:stockdata_crud')  # Adjust redirect as needed
#     else:
#         form = StockDataForm(instance=stock)
#         report_formset = ReportFormSet(instance=stock, prefix='report')
#         shareholding_formset = ShareholdingPatternFormSet(instance=stock, prefix='shareholding')
#         relation_formset = CompanyRelationFormSet(instance=stock, prefix='relation')
#         business_formset = PrincipalBusinessActivityFormSet(instance=stock, prefix='business')

#     return render(request, 'stockData_edit_modal.html', {
#         'form': form,
#         'report_formset': report_formset,
#         'shareholding_formset': shareholding_formset,
#         'relation_formset': relation_formset,
#         'business_formset': business_formset,
#         'title': 'Edit Stock'
#     })


from django.shortcuts import render, get_object_or_404, redirect
from .forms import (
    StockDataForm,
    ReportFormSet,
    ShareholdingPatternFormSet,
    CompanyRelationFormSet,
    PrincipalBusinessActivityFormSet
)

def edit_stockdata(request, pk):
    stock = get_object_or_404(StockData, pk=pk)

    if request.method == 'POST':
        form = StockDataForm(request.POST, request.FILES, instance=stock)
        report_formset = ReportFormSet(request.POST, instance=stock, prefix='report')
        shareholding_formset = ShareholdingPatternFormSet(request.POST, instance=stock, prefix='shareholding')
        relation_formset = CompanyRelationFormSet(request.POST, instance=stock, prefix='relation')
        business_formset = PrincipalBusinessActivityFormSet(request.POST, instance=stock, prefix='business')

        if all([
            form.is_valid(),
            report_formset.is_valid(),
            shareholding_formset.is_valid(),
            relation_formset.is_valid(),
            business_formset.is_valid()
        ]):
            form.save()
            report_formset.save()
            shareholding_formset.save()
            relation_formset.save()
            business_formset.save()
            return redirect('SM_User:stockdata_crud')
        else:
            # Debug output (you can remove these in production)
            print("Form valid?", form.is_valid(), form.errors)
            print("Report formset valid?", report_formset.is_valid(), report_formset.errors)
            print("Shareholding formset valid?", shareholding_formset.is_valid(), shareholding_formset.errors)
            print("Relation formset valid?", relation_formset.is_valid(), relation_formset.errors)
            print("Business formset valid?", business_formset.is_valid(), business_formset.errors)

    else:
        form = StockDataForm(instance=stock)
        report_formset = ReportFormSet(instance=stock, prefix='report')
        shareholding_formset = ShareholdingPatternFormSet(instance=stock, prefix='shareholding')
        relation_formset = CompanyRelationFormSet(instance=stock, prefix='relation')
        business_formset = PrincipalBusinessActivityFormSet(instance=stock, prefix='business')

    return render(request, 'stockData_edit_modal.html', {
        'form': form,
        'report_formset': report_formset,
        'shareholding_formset': shareholding_formset,
        'relation_formset': relation_formset,
        'business_formset': business_formset,
        'title': 'Edit Stock',
    })



# from django.shortcuts import render
# from django.views.decorators.csrf import csrf_exempt
# from docx import Document
# from django.utils.html import format_html
# import re

# @csrf_exempt
# def convert_docx_to_html(request):
#     above_outlook = ""
#     between_outlook_and_shareholding = ""
#     below_shareholding = ""

#     if request.method == 'POST' and request.FILES.get('word_file'):
#         docx_file = request.FILES['word_file']
#         document = Document(docx_file)

#         html_output = ""

#         for para in document.paragraphs:
#             style = para.style.name.lower()
#             line_html = ""

#             # Lists
#             if 'list paragraph' in style:
#                 text = para.text.strip()
#                 if text.startswith('-'):
#                     line_html = f"<ul><li>{text[1:].strip()}</li></ul>"
#                 elif re.match(r'^\d+[\.\)]', text):
#                     line_html = f"<ol><li>{text[2:].strip()}</li></ol>"
#                 else:
#                     line_html = f"<p>{text}</p>"

#             # Headings
#             elif 'heading 1' in style:
#                 line_html = f'<h1>{para.text}</h1>'
#             elif 'heading 2' in style:
#                 line_html = f'<h2>{para.text}</h2>'
#             else:
#                 for run in para.runs:
#                     run_text = run.text
#                     if not run_text:
#                         continue
#                     if run.bold:
#                         run_text = f"<strong>{run_text}</strong>"
#                     if run.italic:
#                         run_text = f"<em>{run_text}</em>"
#                     if run.underline:
#                         run_text = f"<u>{run_text}</u>"
#                     line_html += run_text

#                 if line_html.strip():
#                     line_html = f"<p>{line_html}</p>"

#             html_output += line_html

#         # === SECTION SPLIT LOGIC ===
#         outlook_marker = '<strong>INDUSTRY OUTLOOK</strong></p>'
#         shareholding_marker = 'Latest shareholding pattern'

#         outlook_idx = html_output.find(outlook_marker)
#         shareholding_idx = html_output.find(shareholding_marker)

#         # if outlook_idx != -1:
#         #     outlook_end = outlook_idx + len(outlook_marker)
#         #     above_outlook = html_output[:outlook_end]
#         #     if shareholding_idx != -1:
#         #         between_outlook_and_shareholding = html_output[outlook_end:shareholding_idx]
#         #         below_shareholding = html_output[shareholding_idx:]
#         #     else:
#         #         between_outlook_and_shareholding = html_output[outlook_end:]
#         # else:
#         #     above_outlook = html_output
        
#         if outlook_idx != -1:
#             outlook_break = '<br><br>\n\n' + outlook_marker
#             html_output = html_output.replace(outlook_marker, outlook_break, 1)

#             outlook_idx = html_output.find(outlook_break)
#             outlook_end = outlook_idx + len(outlook_break)

#             above_outlook = html_output[:outlook_end] + '\n'
#             if shareholding_idx != -1:
#                 between_outlook_and_shareholding = html_output[outlook_end:shareholding_idx]
#                 below_shareholding = html_output[shareholding_idx:]
#             else:
#                 between_outlook_and_shareholding = html_output[outlook_end:]
#         else:
#             above_outlook = html_output + '\n'

#         raw_html_combined = (
#             f"{above_outlook}"
#             f"\n\n<!-- === SECTION: INDUSTRY OUTLOOK STARTS HERE === -->\n\n"
#             f"{between_outlook_and_shareholding}"
#             f"\n\n<!-- === SECTION: LATEST SHAREHOLDING STARTS HERE === -->\n\n"
#             f"{below_shareholding}"
#         )

#         context = {
#             'above_outlook': format_html(above_outlook),
#             'between_outlook_and_shareholding': format_html(between_outlook_and_shareholding),
#             'below_shareholding': format_html(below_shareholding),
#             'converted_html': raw_html_combined,
#         }
#         return render(request, 'htmlPreview/upload_and_preview.html', context)

#     return render(request, 'htmlPreview/upload_and_preview.html')




















import re
from docx import Document
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render
from django.http import HttpResponse
from django.utils.html import format_html

# Store latest HTML in-memory for download (for demo only)
last_converted_html = ""

import re
from docx import Document
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render
from django.utils.html import escape

@csrf_exempt
def convert_docx_to_html(request):
    converted_html = ""

    if request.method == 'POST' and request.FILES.get('word_file'):
        document = Document(request.FILES['word_file'])
        html_output = ""

        # Convert paragraphs
        for para in document.paragraphs:
            text = para.text.strip()
            style = para.style.name.lower()

            # Use <h2> for headings
            if 'heading' in style:
                html_output += f"<h2>{escape(text)}</h2>\n"
            else:
                paragraph_content = ""
                for run in para.runs:
                    run_text = escape(run.text)
                    if not run_text.strip():
                        continue

                    # Wrap styled text in <span>
                    if run.bold or run.italic or run.underline:
                        paragraph_content += f"<span>{run_text}</span>"
                    else:
                        paragraph_content += run_text

                if paragraph_content.strip():
                    html_output += f"<p>{paragraph_content}</p>\n"

        # Convert tables
        for table in document.tables:
            html_output += "<table>\n"
            for row in table.rows:
                html_output += "<p>"
                for cell in row.cells:
                    html_output += f"<span>{escape(cell.text.strip())}</span> "
                html_output += "</p>\n"
            html_output += "</table>\n"

        converted_html = html_output

    return render(request, 'htmlPreview/upload_and_preview.html', {
        'converted_html': converted_html
    })





















































from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from unlisted_stock_marketplace.models import CustomFieldDefinition, CustomFieldValue, TableHeader
from decimal import Decimal
import json


def assign_next_order(definition):
    max_order = CustomFieldValue.objects.filter(field_definition=definition).aggregate(Max('order'))['order__max']
    return (max_order or 0) + 1


def shift_order_if_needed(definition, target_order):
    affected = CustomFieldValue.objects.filter(field_definition=definition, order__gte=target_order).order_by('-order')
    for obj in affected:
        obj.order += 1
        obj.save()


def values_for_definition(request, def_id):
    definition = get_object_or_404(CustomFieldDefinition, pk=def_id)
    all_values = CustomFieldValue.objects.filter(field_definition=definition).select_related('table_header', 'parent_field_value').order_by('order')

    grouped = []
    parent_map = {}

    for val in all_values:
        if val.parent_field_value_id is None:
            parent_map[val.pk] = {'parent': val, 'children': []}
        else:
            if val.parent_field_value_id in parent_map:
                parent_map[val.parent_field_value_id]['children'].append(val)

    grouped = list(parent_map.values())
    headers = TableHeader.objects.all().order_by('order')

    return render(request, 'contect/definition_values.html', {
        'definition': definition,
        'grouped': grouped,
        'headers': headers
    })


@csrf_exempt
def add_custom_value(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        def_id = data.get('definition_id')
        definition = CustomFieldDefinition.objects.get(id=def_id)

        value = CustomFieldValue(
            field_definition=definition,
            name=data.get('name'),
            description=data.get('description'),
            table_header_id=data.get('table_header'),
            parent_field_value_id=data.get('parent_id'),
            text_style=data.get('text_style', 'normal')
        )

        val_type = definition.field_type
        if val_type == 'dec':
            value.dec_value = Decimal(data.get('value')) if data.get('value') else None
        elif val_type == 'int':
            value.int_value = int(data.get('value')) if data.get('value') else None
        elif val_type == 'char':
            value.char_value = data.get('value')
        elif val_type == 'date':
            value.date_value = data.get('value')

        value.order = assign_next_order(definition)
        value.save()
        return JsonResponse({'status': 'ok'})


@csrf_exempt
def update_custom_value(request, pk):
    if request.method == 'POST':
        data = json.loads(request.body)
        value = get_object_or_404(CustomFieldValue, pk=pk)

        value.name = data.get('name')
        value.description = data.get('description')
        value.table_header_id = data.get('table_header')
        value.parent_field_value_id = data.get('parent_id')
        value.text_style = data.get('text_style', 'normal')


        val_type = value.field_definition.field_type
        if val_type == 'dec':
            value.dec_value = Decimal(data.get('value')) if data.get('value') else None
        elif val_type == 'int':
            value.int_value = int(data.get('value')) if data.get('value') else None
        elif val_type == 'char':
            value.char_value = data.get('value')
        elif val_type == 'date':
            value.date_value = data.get('value')

        value.save()
        return JsonResponse({'status': 'updated'})


@csrf_exempt
def delete_custom_value(request, pk):
    if request.method == 'POST':
        value = get_object_or_404(CustomFieldValue, pk=pk)
        value.delete()
        return JsonResponse({'status': 'deleted'})
@csrf_exempt
def bulk_upload_values(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            def_id = int(data.get('definition_id'))
            definition = get_object_or_404(CustomFieldDefinition, id=def_id)

            raw = data.get('bulk_data', '')
            header_id = int(data.get('starting_header'))

            table_headers = list(TableHeader.objects.all().order_by('id'))  # order by ID for consistent alignment

            try:
                start_index = [th.id for th in table_headers].index(header_id)
            except ValueError:
                return JsonResponse({'status': 'invalid_header'})

            lines = [line.strip() for line in raw.strip().splitlines() if line.strip()]

            for line in lines:
                parts = line.split()
                name_parts = [x for x in parts if not is_number(x)]
                val_parts = [x for x in parts if is_number(x)]

                name = ' '.join(name_parts)
                try:
                    values = [Decimal(v) for v in val_parts]
                except Exception:
                    continue  # skip invalid line

                headers = table_headers[start_index:start_index + len(values)]
                if len(headers) < len(values) or not values or not name:
                    continue  # skip invalid rows

                # Create parent
                parent = CustomFieldValue.objects.create(
                    field_definition=definition,
                    name=name,
                    dec_value=values[0] if definition.field_type == 'dec' else None,
                    table_header=headers[0],
                    order=assign_next_order(definition),
                    text_style=data.get('text_style', 'normal')
                )

                # Create children
                for idx in range(1, len(values)):
                    CustomFieldValue.objects.create(
                        field_definition=definition,
                        name=None,
                        dec_value=values[idx] if definition.field_type == 'dec' else None,
                        table_header=headers[idx],
                        parent_field_value=parent,
                        order=assign_next_order(definition)
                    )

            return JsonResponse({'status': 'bulk_uploaded'})

        except Exception as e:
            return JsonResponse({'status': f'error: {str(e)}'})



@csrf_exempt
def reorder_custom_values(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        order_list = data.get('order')

        for idx, val_id in enumerate(order_list):
            CustomFieldValue.objects.filter(id=val_id).update(order=idx + 1)

        return JsonResponse({'status': 'reordered'})


def is_number(s):
    try:
        float(s)
        return True
    except Exception:
        return False


# views.py
from django.shortcuts import render
from unlisted_stock_marketplace.models import *

def stock_model_overview(request):
    selected_stock_id = request.GET.get("stock")
    grouped_data = {}
    stocks = StockData.objects.all().order_by("company_name")

    # Filter by selected stock if present
    definitions = CustomFieldDefinition.objects.select_related('stock').order_by('stock__company_name', 'model_type')
    if selected_stock_id:
        definitions = definitions.filter(stock_id=selected_stock_id)

    for defn in definitions:
        stock = defn.stock
        if stock not in grouped_data:
            grouped_data[stock] = []
        grouped_data[stock].append(defn)

    return render(request, 'contect/stock_model_overview.html', {
        'grouped_data': grouped_data,
        'stocks': stocks,
        'selected_stock_id': int(selected_stock_id) if selected_stock_id else None
    })


# app/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse, reverse_lazy
from django.core.paginator import Paginator
from django.utils.dateparse import parse_date
from django.contrib.auth.decorators import login_required

from .forms import StockHistoryForm


from django.db.models import Count

# app/views.py

from django.core.paginator import Paginator
from django.contrib.auth.decorators import login_required
from django.db.models import F, Count, Window


@login_required
def stockhistory_list(request, stock_id=None):
    # Always show latest updated/added first
    qs = StockHistory.objects.select_related("stock").order_by("-updated_at", "-id")

    # Filters
    stock_id = stock_id or request.GET.get("stock_id")
    date_exact = request.GET.get("date")
    q = request.GET.get("q")

    if stock_id:
        qs = qs.filter(stock_id=stock_id)

    if date_exact:
        try:
            qs = qs.filter(timestamp=datetime.strptime(date_exact, "%Y-%m-%d").date())
        except ValueError:
            pass

    if q:
        qs = qs.filter(stock__company_name__icontains=q)

    # Annotate duplicate count (same stock + date)
    qs = qs.annotate(
        duplicate_count=Window(
            expression=Count("id"),
            partition_by=[F("stock_id"), F("timestamp")],
        )
    )

    latest_entry = qs.first()

    # Build grouped IDs
    group_ids = {}
    for row in qs.values("id", "stock_id", "timestamp"):
        key = f"{row['stock_id']}|{row['timestamp']}"
        group_ids.setdefault(key, []).append(row["id"])

    # Pagination
    paginator = Paginator(qs, 20)
    page_obj = paginator.get_page(request.GET.get("page"))

    # Attach group_ids to page items
    for item in page_obj.object_list:
        key = f"{item.stock_id}|{item.timestamp}"
        ids_list = group_ids.get(key, [])
        item.group_ids = sorted(ids_list, reverse=True)
        item.group_ids_str = ", ".join(str(i) for i in item.group_ids)

    params = request.GET.copy()
    params.pop("page", None)
    querystring = params.urlencode()

    current_stock = StockData.objects.filter(id=stock_id).first() if stock_id else None

    return render(request, "stockhistory/history_list.html", {
        "histories": page_obj,
        "all_stocks": StockData.objects.all(),
        "current_stock": current_stock,
        "is_paginated": page_obj.has_other_pages(),
        "page_obj": page_obj,
        "querystring": querystring,
        "latest_entry": latest_entry,
    })



@login_required
def stockhistory_detail(request, pk):
    history = get_object_or_404(StockHistory, pk=pk)
    return render(request, "stockhistory/history_detail.html", {"history": history})


@login_required
def stockhistory_create(request, stock_id=None):
    if request.method == "POST":
        form = StockHistoryForm(request.POST)
        if form.is_valid():
            history = form.save()
            return redirect("SM_User:Historylist")
    else:
        initial = {}
        if stock_id:
            initial["stock"] = get_object_or_404(StockData, id=stock_id)
        form = StockHistoryForm(initial=initial)

    return render(request, "stockhistory/history_form.html", {"form": form})


@login_required
def stockhistory_update(request, pk):
    history = get_object_or_404(StockHistory, pk=pk)
    if request.method == "POST":
        form = StockHistoryForm(request.POST, instance=history)
        if form.is_valid():
            form.save()
            return redirect("SM_User:Historydetail", pk=history.pk)
    else:
        form = StockHistoryForm(instance=history)

    return render(request, "stockhistory/history_form.html", {"form": form})


@login_required
def stockhistory_delete(request, pk):
    history = get_object_or_404(StockHistory, pk=pk)
    if request.method == "POST":
        stock_id = history.stock_id
        history.delete()
        if stock_id:
            return redirect("SM_User:Historylist")
        return redirect("SM_User:Historylist")

    return render(request, "stockhistory/history_confirm_delete.html", {"object": history})




# CRUD Events 
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import Event
from .forms import EventForm

# List all visible events
def event_list(request):
    events = Event.objects.filter(show=True).order_by("-date_time")
    return render(request, "events/event_list.html", {"events": events})

# Single event detail
def event_detail(request, pk):
    event = get_object_or_404(Event, pk=pk, show=True)
    return render(request, "events/event_detail.html", {"event": event})

# Create event
@login_required
def event_create(request):
    if request.method == "POST":
        form = EventForm(request.POST, request.FILES)
        if form.is_valid():
            event = form.save(commit=False)
            event.user = request.user
            event.save()
            return redirect("SM_User:event_list")
    else:
        form = EventForm()
    return render(request, "events/event_form.html", {"form": form, "title": "Create Event"})

# Update event
@login_required
def event_update(request, pk):
    event = get_object_or_404(Event, pk=pk, user=request.user)
    if request.method == "POST":
        form = EventForm(request.POST, request.FILES, instance=event)
        if form.is_valid():
            form.save()
            return redirect("SM_User:event_detail", pk=pk)
    else:
        form = EventForm(instance=event)
    return render(request, "events/event_form.html", {"form": form, "title": "Update Event"})

# Delete event
@login_required
def event_delete(request, pk):
    event = get_object_or_404(Event, pk=pk, user=request.user)
    if request.method == "POST":
        event.delete()
        return redirect("SM_User:event_list")
    return render(request, "events/event_confirm_delete.html", {"event": event})

# views.py
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML
import tempfile
from .models import Event  # update model import
# views.py
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML
import tempfile
from .models import Event  # update model import

# views.py
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.conf import settings
from weasyprint import HTML
from .models import Event  # adjust import if needed

from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404
import os

def event_pdf(request, pk):
    event = get_object_or_404(Event, pk=pk)

    # Check if document exists
    if not event.document:
        raise Http404("No document uploaded for this event.")

    # Use Django's FileResponse to serve the PDF
    file_path = event.document.path
    if not os.path.exists(file_path):
        raise Http404("Document file not found.")

    response = FileResponse(open(file_path, "rb"), content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{os.path.basename(file_path)}"'
    return response


